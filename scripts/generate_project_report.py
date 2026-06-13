from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PilImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report_assets"
OUT = DOCS / "CEN306_My_Notebook_Project_Report.docx"
PDF_OUT = DOCS / "CEN306_My_Notebook_Project_Report.pdf"


def set_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run(prefix, bold=True)
        body = paragraph.add_run(text[len(bold_prefix) :])
        set_run(body)
    else:
        run = paragraph.add_run(text)
        set_run(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run(run)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run(run)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string("2E74B5" if level < 3 else "1F4D78")
    return paragraph


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.7)
    table.columns[1].width = Inches(4.8)
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Details"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run(run, bold=True)

    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    doc.add_paragraph()
    return table


def add_screenshot(doc, filename, caption):
    path = ASSETS / filename
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(2.35))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    set_run(caption_run, size=9, color="555555")


def build_report():
    DOCS.mkdir(exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("My Notebook Flutter Mobile Application")
    set_run(title_run, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("CEN306 Final Project Report")
    set_run(subtitle_run, size=13, color="555555")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("Prepared for Mobile Application Design and Development")
    set_run(meta_run, size=10, color="555555")

    add_heading(doc, "1. Project Purpose and Scope", 1)
    add_paragraph(
        doc,
        "My Notebook is a Flutter study notebook application designed to help students organize subjects, notes, tasks, reminders, search results, and study resources in one local-first app. The application supports different study modes such as Basic, Business, Engineering, and Medicine, then adjusts resources and dashboard content based on the selected field.",
    )
    add_paragraph(
        doc,
        "The main goal is to provide a practical student workspace rather than a simple counter or sample app. Users can register, log in, select a field, create subjects with visual notebook covers, write or draw notes, add sticky notes and emoji stickers, manage tasks and reminders, search stored content, use field-aware resources, and customize the app theme.",
    )

    add_heading(doc, "2. Main Functional Screens", 1)
    add_key_value_table(
        doc,
        [
            ("Authentication", "Login and sign-up screens register users and verify hashed passwords."),
            ("Mode Selection", "Field selection chooses the study context used throughout the app."),
            ("Dashboard", "Home screen summarizes subjects, notes, tasks, and reminders for the selected field."),
            ("Subjects", "Users create, edit, delete, and open subjects with custom notebook covers."),
            ("Note Editor", "Users type, draw, erase, pick pen types, add sticky notes, add emoji stickers, and save notes."),
            ("Tasks and Reminders", "Users track study tasks and dated reminders scoped to the current field and user."),
            ("Search", "Users search subjects, notes, tasks, and reminders from one screen."),
            ("Resources", "Field-aware study templates and prompts can become subjects or tasks."),
            ("Theme", "Users select a theme color that updates the Material color scheme."),
        ],
    )

    add_heading(doc, "3. Screen Evidence", 1)
    add_paragraph(
        doc,
        "The screenshots below show the current mobile web rendering used during verification. The same Flutter UI is shared with mobile builds.",
    )
    add_screenshot(doc, "01_mobile_home.png", "Figure 1. Mobile dashboard with responsive summary cards.")
    add_screenshot(doc, "02_mobile_drawer.png", "Figure 2. Drawer navigation on a compact mobile viewport.")
    add_screenshot(doc, "04_subject_cover.png", "Figure 3. Subject list with visual notebook cover.")
    add_screenshot(doc, "05_note_editor.png", "Figure 4. Note editor with pen-type tray and drawing tools.")
    add_screenshot(doc, "06_note_sticky.png", "Figure 5. Sticky note attached to the notebook page.")

    add_heading(doc, "4. Layered Architecture", 1)
    add_paragraph(
        doc,
        "The project is organized into separate UI, controller/business, repository, DAO, database, model, utility, and widget layers. This separation reduces direct coupling between Flutter screens and database implementation details.",
    )
    add_bullet(doc, "UI layer: lib/screens and lib/widgets contain visual layouts, navigation, forms, dialogs, and reusable UI elements.")
    add_bullet(doc, "Controller layer: lib/controllers exposes actions used by screens, such as login, loading notes, adding tasks, and searching.")
    add_bullet(doc, "Repository layer: lib/repositories applies current-user scoping and provides app-facing data operations.")
    add_bullet(doc, "DAO layer: lib/dao contains low-level table access for users, subjects, notes, tasks, and reminders.")
    add_bullet(doc, "Database layer: lib/database defines the database adapter interface plus SQLite and web implementations.")

    add_heading(doc, "5. SQLite, DAO, Repository, and Data Flow", 1)
    add_paragraph(
        doc,
        "On mobile and desktop targets the app uses sqflite through DBHelper. The web adapter mirrors the same AppDatabase interface with SharedPreferences so the app can still be demonstrated in Chrome. The core mobile data layer is SQLite.",
    )
    add_key_value_table(
        doc,
        [
            ("users", "id, name, email, password. Stores account records with PBKDF2 password hashes."),
            ("subjects", "id, title, field, userEmail, coverColor, coverPattern. Stores user-scoped subject notebooks."),
            ("notes", "id, subject, field, content, noteType, drawing, userEmail. Stores typed text, drawing JSON, sticky notes, stickers, and note type."),
            ("tasks", "id, title, isCompleted, field, userEmail. Stores user-scoped task checklist records."),
            ("reminders", "id, title, reminderDate, field, userEmail. Stores dated reminder records."),
        ],
    )
    add_paragraph(doc, "Example note save flow:")
    add_number(doc, "SubjectNoteScreen collects text, strokes, sticky notes, stickers, note type, subject, and field.")
    add_number(doc, "NoteController receives the save request from the UI.")
    add_number(doc, "NoteRepository adds the current user email scope.")
    add_number(doc, "NoteDao performs the upsert operation against the notes table.")
    add_number(doc, "DBHelper supplies the platform database adapter, SQLite on mobile and a web adapter for Chrome demos.")

    add_heading(doc, "6. State Management", 1)
    add_paragraph(
        doc,
        "The app uses StatefulWidget and setState for screens that hold changing data, including home counts, subjects, notes, tasks, reminders, search results, resource category selection, and the note editor. Theme color uses a ValueNotifier so MaterialApp can rebuild when the user selects a new theme.",
    )
    add_bullet(doc, "SubjectNoteScreen manages active tool, pen type, stroke list, eraser preview, sticky notes, stickers, selected template, and text editing state.")
    add_bullet(doc, "HomeScreen updates dashboard counts after navigation and data changes.")
    add_bullet(doc, "TasksScreen and RemindersScreen reload and update lists after CRUD actions.")
    add_bullet(doc, "ThemeScreen uses AppThemeController.seedColor to update the app color scheme persistently.")

    add_heading(doc, "7. Testing and Build Verification", 1)
    add_key_value_table(
        doc,
        [
            ("dart format lib test", "Passed with no files needing changes on the final run."),
            ("flutter analyze", "Passed with no issues found."),
            ("flutter test", "Passed all 6 tests."),
            ("flutter build web", "Passed and generated build/web."),
            ("flutter build apk --debug", "Passed and generated app-debug.apk."),
            ("Visual check", "Mobile web dashboard and drawer were checked at 390 x 844 viewport. Overflow was found, fixed, and rechecked."),
        ],
    )

    add_heading(doc, "8. Development Challenges and Solutions", 1)
    add_bullet(doc, "Startup session handling originally sent logged-in users back to login; the startup flow now restores saved sessions and fields.")
    add_bullet(doc, "Plain-text password storage was replaced with PBKDF2 SHA-256 hashing while still supporting legacy password upgrade on login.")
    add_bullet(doc, "Field-only data separation was strengthened with userEmail scoping across notes, subjects, tasks, and reminders.")
    add_bullet(doc, "The note editor grew from a basic note page into a visual notebook with pen types, improved eraser behavior, sticky notes, and emoji stickers.")
    add_bullet(doc, "The desktop-style sidebar did not suit mobile widths; HomeScreen now switches to drawer navigation on compact screens.")
    add_bullet(doc, "Android builds currently pass but show a future Kotlin Gradle Plugin warning from shared_preferences_android; this should be resolved before app-store release by upgrading the plugin when a compatible version is available.")

    add_heading(doc, "9. Conclusion", 1)
    add_paragraph(
        doc,
        "The current project now satisfies the main code-side expectations of the rubric: meaningful screens, working navigation, stateful UI behavior, SQLite-backed persistence, DAO/repository data flow, and successful build/test verification. The remaining work before final submission is to review this report, add student/course identity details required by the instructor, and optionally add more screenshots if the final demonstration needs every screen shown individually.",
    )

    doc.save(OUT)
    print(OUT)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ReportTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=21,
            leading=26,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportSubtitle",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=12,
            leading=16,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#555555"),
            spaceAfter=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionHeading",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Caption",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#555555"),
            spaceAfter=12,
        )
    )
    return styles


def pdf_table(rows, styles, col_widths=(1.65 * inch, 4.35 * inch)):
    data = [
        [
            Paragraph("<b>Item</b>", styles["Body"]),
            Paragraph("<b>Details</b>", styles["Body"]),
        ]
    ]
    for key, value in rows:
        data.append([Paragraph(key, styles["Body"]), Paragraph(value, styles["Body"])])
    table = Table(data, colWidths=col_widths, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#DADCE0")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def pdf_bullets(items, styles):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["Body"])) for item in items],
        bulletType="bullet",
        leftIndent=18,
        bulletFontName="Helvetica",
        bulletFontSize=8,
    )


def pdf_numbers(items, styles):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["Body"])) for item in items],
        bulletType="1",
        leftIndent=18,
        bulletFontName="Helvetica",
        bulletFontSize=9,
    )


def pdf_screenshot(filename, caption, styles):
    path = ASSETS / filename
    if not path.exists():
        return []
    with PilImage.open(path) as image:
        width, height = image.size
    target_width = 2.15 * inch
    target_height = target_width * height / width
    return [
        Image(str(path), width=target_width, height=target_height, hAlign="CENTER"),
        Paragraph(caption, styles["Caption"]),
    ]


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    story = [
        Paragraph("My Notebook Flutter Mobile Application", styles["ReportTitle"]),
        Paragraph("CEN306 Final Project Report", styles["ReportSubtitle"]),
        Paragraph("1. Project Purpose and Scope", styles["SectionHeading"]),
        Paragraph(
            "My Notebook is a Flutter study notebook application designed to help students organize subjects, notes, tasks, reminders, search results, and study resources in one local-first app. The application supports different study modes such as Basic, Business, Engineering, and Medicine, then adjusts resources and dashboard content based on the selected field.",
            styles["Body"],
        ),
        Paragraph(
            "The app lets users register, log in, select a field, create subjects with visual notebook covers, write or draw notes, add sticky notes and emoji stickers, manage tasks and reminders, search stored content, use field-aware resources, and customize the app theme.",
            styles["Body"],
        ),
        Paragraph("2. Main Functional Screens", styles["SectionHeading"]),
        pdf_table(
            [
                ("Authentication", "Login and sign-up screens register users and verify hashed passwords."),
                ("Mode Selection", "Field selection chooses the study context used throughout the app."),
                ("Dashboard", "Home screen summarizes subjects, notes, tasks, and reminders for the selected field."),
                ("Subjects", "Users create, edit, delete, and open subjects with custom notebook covers."),
                ("Note Editor", "Users type, draw, erase, pick pen types, add sticky notes, add emoji stickers, and save notes."),
                ("Tasks and Reminders", "Users track study tasks and dated reminders scoped to the current field and user."),
                ("Search", "Users search subjects, notes, tasks, and reminders from one screen."),
                ("Resources", "Field-aware study templates and prompts can become subjects or tasks."),
                ("Theme", "Users select a theme color that updates the Material color scheme."),
            ],
            styles,
        ),
        Spacer(1, 10),
        Paragraph("3. Screen Evidence", styles["SectionHeading"]),
        Paragraph(
            "The screenshots below show the current mobile web rendering used during verification. The same Flutter UI is shared with mobile builds.",
            styles["Body"],
        ),
    ]
    for filename, caption in [
        ("01_mobile_home.png", "Figure 1. Mobile dashboard with responsive summary cards."),
        ("02_mobile_drawer.png", "Figure 2. Drawer navigation on a compact mobile viewport."),
        ("04_subject_cover.png", "Figure 3. Subject list with visual notebook cover."),
        ("05_note_editor.png", "Figure 4. Note editor with pen-type tray and drawing tools."),
        ("06_note_sticky.png", "Figure 5. Sticky note attached to the notebook page."),
    ]:
        story.extend(pdf_screenshot(filename, caption, styles))

    story.extend(
        [
            Paragraph("4. Layered Architecture", styles["SectionHeading"]),
            Paragraph(
                "The project is organized into separate UI, controller/business, repository, DAO, database, model, utility, and widget layers. This separation reduces direct coupling between Flutter screens and database implementation details.",
                styles["Body"],
            ),
            pdf_bullets(
                [
                    "UI layer: lib/screens and lib/widgets contain visual layouts, navigation, forms, dialogs, and reusable UI elements.",
                    "Controller layer: lib/controllers exposes actions used by screens.",
                    "Repository layer: lib/repositories applies current-user scoping and app-facing data operations.",
                    "DAO layer: lib/dao contains low-level table access for users, subjects, notes, tasks, and reminders.",
                    "Database layer: lib/database defines the database adapter interface plus SQLite and web implementations.",
                ],
                styles,
            ),
            Paragraph("5. SQLite, DAO, Repository, and Data Flow", styles["SectionHeading"]),
            Paragraph(
                "On mobile and desktop targets the app uses sqflite through DBHelper. The web adapter mirrors the same AppDatabase interface with SharedPreferences so the app can still be demonstrated in Chrome. The core mobile data layer is SQLite.",
                styles["Body"],
            ),
            pdf_table(
                [
                    ("users", "id, name, email, password. Stores account records with PBKDF2 password hashes."),
                    ("subjects", "id, title, field, userEmail, coverColor, coverPattern. Stores user-scoped subject notebooks."),
                    ("notes", "id, subject, field, content, noteType, drawing, userEmail. Stores typed text, drawing JSON, sticky notes, stickers, and note type."),
                    ("tasks", "id, title, isCompleted, field, userEmail. Stores user-scoped task checklist records."),
                    ("reminders", "id, title, reminderDate, field, userEmail. Stores dated reminder records."),
                ],
                styles,
            ),
            Spacer(1, 8),
            Paragraph("Example note save flow:", styles["Body"]),
            pdf_numbers(
                [
                    "SubjectNoteScreen collects text, strokes, sticky notes, stickers, note type, subject, and field.",
                    "NoteController receives the save request from the UI.",
                    "NoteRepository adds the current user email scope.",
                    "NoteDao performs the upsert operation against the notes table.",
                    "DBHelper supplies the platform database adapter.",
                ],
                styles,
            ),
            Paragraph("6. State Management", styles["SectionHeading"]),
            Paragraph(
                "The app uses StatefulWidget and setState for screens that hold changing data, including home counts, subjects, notes, tasks, reminders, search results, resource category selection, and the note editor. Theme color uses a ValueNotifier so MaterialApp can rebuild when the user selects a new theme.",
                styles["Body"],
            ),
            pdf_bullets(
                [
                    "SubjectNoteScreen manages active tool, pen type, stroke list, eraser preview, sticky notes, stickers, selected template, and text editing state.",
                    "HomeScreen updates dashboard counts after navigation and data changes.",
                    "TasksScreen and RemindersScreen reload and update lists after CRUD actions.",
                    "ThemeScreen uses AppThemeController.seedColor to update the app color scheme persistently.",
                ],
                styles,
            ),
            Paragraph("7. Testing and Build Verification", styles["SectionHeading"]),
            pdf_table(
                [
                    ("dart format lib test", "Passed with no files needing changes on the final run."),
                    ("flutter analyze", "Passed with no issues found."),
                    ("flutter test", "Passed all 6 tests."),
                    ("flutter build web", "Passed and generated build/web."),
                    ("flutter build apk --debug", "Passed and generated app-debug.apk."),
                    ("Visual check", "Mobile web dashboard and drawer were checked at 390 x 844 viewport. Overflow was found, fixed, and rechecked."),
                ],
                styles,
            ),
            Spacer(1, 8),
            Paragraph("8. Development Challenges and Solutions", styles["SectionHeading"]),
            pdf_bullets(
                [
                    "Startup session handling now restores saved sessions and fields.",
                    "Plain-text password storage was replaced with PBKDF2 SHA-256 hashing while supporting legacy upgrade on login.",
                    "Field-only data separation was strengthened with userEmail scoping across subjects, notes, tasks, and reminders.",
                    "The note editor now includes pen types, improved eraser behavior, sticky notes, and emoji stickers.",
                    "HomeScreen now switches to drawer navigation on compact screens.",
                    "Android builds currently pass but show a future Kotlin Gradle Plugin warning from shared_preferences_android.",
                ],
                styles,
            ),
            Paragraph("9. Conclusion", styles["SectionHeading"]),
            Paragraph(
                "The current project satisfies the main code-side expectations of the rubric: meaningful screens, working navigation, stateful UI behavior, SQLite-backed persistence, DAO/repository data flow, and successful build/test verification. Before final submission, student/course identity details can be added if required by the instructor.",
                styles["Body"],
            ),
        ]
    )

    doc.build(story)
    print(PDF_OUT)


if __name__ == "__main__":
    build_report()
    build_pdf()
