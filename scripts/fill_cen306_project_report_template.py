from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path("/Users/favourokwudili/Downloads/CEN306_Project_Report_Template (2).docx")
DOCS = ROOT / "docs"
ASSETS = DOCS / "report_assets"
OUT = DOCS / "CEN306_My_Notebook_Filled_Project_Report.docx"
BULLET_NUM_ID = None
NUMBER_NUM_ID = None


def clear_document(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_style(run, size=11, bold=False, color="000000", font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, fill=None, color="000000", size=10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_style(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)


def _max_numbering_id(numbering, tag_name, attr_name, default):
    values = []
    for element in numbering.findall(qn(tag_name)):
        value = element.get(qn(attr_name))
        if value is not None and value.isdigit():
            values.append(int(value))
    return max(values, default=default)


def _append_child(parent, tag_name, attrs=None):
    child = OxmlElement(tag_name)
    for key, value in (attrs or {}).items():
        child.set(qn(key), str(value))
    parent.append(child)
    return child


def _create_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_id = _max_numbering_id(
        numbering,
        "w:abstractNum",
        "w:abstractNumId",
        -1,
    ) + 1
    num_id = _max_numbering_id(numbering, "w:num", "w:numId", 0) + 1

    abstract = _append_child(
        numbering,
        "w:abstractNum",
        {"w:abstractNumId": abstract_id},
    )
    _append_child(abstract, "w:multiLevelType", {"w:val": "singleLevel"})

    level = _append_child(abstract, "w:lvl", {"w:ilvl": 0})
    _append_child(level, "w:start", {"w:val": 1})
    _append_child(
        level,
        "w:numFmt",
        {"w:val": "bullet" if kind == "bullet" else "decimal"},
    )
    _append_child(
        level,
        "w:lvlText",
        {"w:val": "\\u2022" if kind == "bullet" else "%1."},
    )
    _append_child(level, "w:lvlJc", {"w:val": "left"})

    p_pr = _append_child(level, "w:pPr")
    _append_child(p_pr, "w:ind", {"w:left": 720, "w:hanging": 360})

    num = _append_child(numbering, "w:num", {"w:numId": num_id})
    _append_child(num, "w:abstractNumId", {"w:val": abstract_id})
    return num_id


def ensure_numbering(doc):
    global BULLET_NUM_ID, NUMBER_NUM_ID
    BULLET_NUM_ID = _create_numbering_definition(doc, "bullet")
    NUMBER_NUM_ID = _create_numbering_definition(doc, "number")


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")

    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def add_paragraph(doc, text="", bold_prefix=None, keep_with_next=False):
    paragraph = doc.add_paragraph()
    if keep_with_next:
        paragraph.paragraph_format.keep_with_next = True

    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_style(prefix, bold=True)
        body = paragraph.add_run(text[len(bold_prefix) :])
        set_run_style(body)
    else:
        run = paragraph.add_run(text)
        set_run_style(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, BULLET_NUM_ID)
    run = paragraph.add_run(text)
    set_run_style(run)


def add_number(doc, text):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, NUMBER_NUM_ID)
    run = paragraph.add_run(text)
    set_run_style(run)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        set_run_style(
            run,
            size=16 if level == 1 else 13 if level == 2 else 11,
            bold=True,
            color="17365D" if level == 1 else "1F4E79",
        )
    return paragraph


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.autofit = True

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, fill="D9EAF7", color="17365D")
        if widths:
            cell.width = Inches(widths[index])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value), size=9.5)
            if widths:
                cells[index].width = Inches(widths[index])

    doc.add_paragraph()
    return table


def add_key_table(doc, rows):
    return add_table(doc, ["Item", "Details"], rows, widths=[1.8, 4.9])


def add_code_block(doc, title, file_name, code, explanation):
    add_heading(doc, title, 3)
    add_paragraph(doc, f"File: {file_name}", bold_prefix="File:")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    set_run_style(run, size=8.5, font="Courier New", color="333333")
    add_paragraph(doc, explanation)


def add_image_cell(cell, image_name, width=1.25):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = ASSETS / image_name
    if path.exists():
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(width))
    else:
        run = paragraph.add_run(f"Screenshot unavailable: {image_name}")
        set_run_style(run, size=9, color="AA0000")


def add_screen_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    headers = [
        "Screen name",
        "Screenshot",
        "Purpose / main function",
        "Navigation / user actions",
    ]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, fill="D9EAF7")

    for screen, image, purpose, actions in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], screen, bold=True)
        add_image_cell(cells[1], image)
        set_cell_text(cells[2], purpose)
        set_cell_text(cells[3], actions)

    doc.add_paragraph()


def build_report():
    DOCS.mkdir(exist_ok=True)
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    ensure_numbering(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CEN306-Mobile Application Design and Development")
    set_run_style(run, size=18, bold=True, color="17365D")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Final Exam Project Report")
    set_run_style(run, size=15, bold=True, color="1F4E79")

    cover_rows = [
        ("Project Title", "My Notebook: Flutter Study and Note-Taking Application"),
        ("Student Name-Surname", "Favour Okwudili - please confirm official spelling"),
        ("Student Number", "TO FILL: Student number"),
        ("Student Signature", "TO FILL: Signature"),
        ("Course Name", "CEN306 - Mobile Application Design and Development"),
        ("Instructor", "Dr. Yildiz Karadayi"),
        ("Submission Date", "13/06/2026 - update if submitting on another date"),
        ("GitHub / Source Code Link", "TO FILL: GitHub/GitLab/Drive source code link"),
        ("Demo Video YouTube Link", "TO FILL: YouTube demo video link, if required"),
        ("Application Name", "My Notebook"),
    ]
    add_table(doc, ["Field", "Completed information"], cover_rows, widths=[2.1, 4.8])

    add_heading(doc, "Academic Integrity Statement", 1)
    add_paragraph(
        doc,
        "I confirm that this report and project were prepared by me and that all external sources, libraries, images, code snippets, packages, and learning resources are properly acknowledged in the References section.",
    )

    add_heading(doc, "Evaluation Alignment Checklist: 30 Points", 1)
    add_table(
        doc,
        ["Official report criterion", "Points", "Where it is addressed", "Evidence in this report"],
        [
            (
                "Report Structure and Organization",
                "5",
                "Cover page, table of contents, all main sections",
                "The report follows the CEN306 template structure and removes generic guidance text.",
            ),
            (
                "Project Purpose and Scope",
                "5",
                "Sections 1-3",
                "Purpose, target users, scope, features, functional requirements, and use cases are explained.",
            ),
            (
                "Architecture and Technical Design",
                "10",
                "Sections 5-7",
                "Layered architecture, controllers, repositories, DAO classes, SQLite/web database adapters, state flow, and code snippets are documented.",
            ),
            (
                "Screenshots and Screen Explanations",
                "5",
                "Section 4 and Appendix A",
                "Main screens are shown with captions, purposes, and user actions.",
            ),
            (
                "Technical Reflection and Process",
                "5",
                "Sections 8-10",
                "Testing, challenges, design decisions, limitations, and future improvements are summarized.",
            ),
        ],
        widths=[2.2, 0.6, 1.8, 2.5],
    )

    add_heading(doc, "Table of Contents", 1)
    for item in [
        "1. Executive Summary",
        "2. Project Introduction",
        "3. Requirements and Use Cases",
        "4. User Interface Design and Screens",
        "5. Technical Architecture",
        "6. Data Layer: SQLite, DAO, and Repository",
        "7. Implementation Details",
        "8. Testing and Quality Assurance",
        "9. Challenges, Design Decisions, and Solutions",
        "10. Conclusion and Future Work",
        "11. References",
        "12. Appendices",
    ]:
        add_paragraph(doc, item)

    add_heading(doc, "1. Executive Summary", 1)
    add_paragraph(
        doc,
        "My Notebook is a Flutter study and note-taking application developed for CEN306 Mobile Application Design and Development. The application helps students organize academic work by combining subject notebooks, visual covers, note pages, drawing tools, sticky notes, emoji stickers, tasks, reminders, search, study resources, theme selection, and local backup export in one app. The target users are students who need a practical local-first workspace for coursework, exam revision, and daily planning. The application is built with Flutter and Dart and uses a layered architecture: screens and widgets handle the interface, controllers handle user actions, repositories provide app-facing data operations, DAO classes isolate database access, and database adapters persist data locally. On mobile and desktop, structured records are stored with SQLite through sqflite. On web, a SharedPreferences-backed adapter implements the same database interface so the app can be tested in Chrome. The project also includes user-scoped data, password hashing, dashboard summaries, responsive navigation, and automated tests for core storage and startup behavior. The main development outcome is a functional student notebook system that goes beyond a basic note app by combining organization, visual note editing, study workflow tools, and a maintainable technical structure.",
    )
    add_key_table(
        doc,
        [
            ("Application name", "My Notebook"),
            ("Problem addressed", "Students need one organized place to manage subjects, notes, study tasks, reminders, and resources."),
            ("Target users", "University students and learners who want a local-first study notebook."),
            ("Core technologies", "Flutter, Dart, SQLite, sqflite, shared_preferences, Material 3."),
            ("Main contribution", "A structured study notebook with visual covers, drawing tools, sticky notes, resources, backup, and layered data architecture."),
        ],
    )

    add_heading(doc, "2. Project Introduction", 1)
    add_heading(doc, "2.1 Project Purpose", 2)
    add_paragraph(
        doc,
        "The purpose of My Notebook is to provide students with a focused digital study workspace. Instead of separating notes, tasks, reminders, and resources across different tools, the application combines them around academic subjects. The app supports common student workflows such as creating a subject, choosing a cover, opening a note page, writing or drawing ideas, adding sticky reminders on the page, tracking tasks, and checking study resources.",
    )
    add_heading(doc, "2.2 Scope of the Project", 2)
    add_paragraph(
        doc,
        "The current version includes authentication, field selection, dashboard summaries, subject management, note creation, typed notes, drawing strokes, multiple pen styles, erasing, sticky notes, emoji stickers, tasks, reminders, search, resources, theme color selection, and local backup export. It is designed for local persistence rather than cloud synchronization. App-store preparation, production signing, cloud backup, and multi-device synchronization are intentionally left as future work after the school submission.",
    )
    add_heading(doc, "2.3 Target Users and Usage Scenario", 2)
    add_paragraph(
        doc,
        "The intended user is a student managing coursework. For example, a Computer Engineering student logs in, selects the Computer Engineering mode, creates an Algorithms subject with a custom notebook cover, opens a study note page, writes lecture points, draws a diagram, attaches a sticky note for revision, adds a task to revise sorting algorithms, sets a reminder for project submission, and later searches the notebook for a keyword.",
    )
    add_heading(doc, "2.4 Main Features", 2)
    add_table(
        doc,
        ["Feature", "Description", "Related screen(s)"],
        [
            ("Authentication", "Users can sign up, log in, log out, and restore a saved session.", "Login, Sign Up, Main startup flow"),
            ("Field selection", "Users select a study field such as Basic, Engineering, Business, or Medicine.", "Field and specialization selection screens"),
            ("Subject notebooks", "Users create, edit, delete, and open subjects with custom visual covers.", "Subjects"),
            ("Visual note editor", "Users type, draw, erase, select pen types, add sticky notes, and add emoji stickers.", "Subject Note"),
            ("Tasks and reminders", "Users create and manage study tasks and dated reminders.", "Tasks, Reminders"),
            ("Search", "Users search across subjects, notes, tasks, and reminders.", "Search"),
            ("Resources", "Users browse field-aware resources and convert some resources into subjects or tasks.", "Resources"),
            ("Theme and backup", "Users select theme color and export their field data as JSON.", "Theme, Backup"),
        ],
        widths=[1.4, 3.6, 1.8],
    )

    add_heading(doc, "3. Requirements and Use Cases", 1)
    add_heading(doc, "3.1 Functional Requirements", 2)
    add_table(
        doc,
        ["ID", "Functional requirement", "Priority", "Implemented?"],
        [
            ("FR-01", "The user can create an account, log in, and log out.", "High", "Yes"),
            ("FR-02", "The user can select and save a study field or specialization.", "High", "Yes"),
            ("FR-03", "The user can create, view, update, and delete subjects.", "High", "Yes"),
            ("FR-04", "The user can choose subject cover colors and patterns.", "Medium", "Yes"),
            ("FR-05", "The user can create and save note content for a subject.", "High", "Yes"),
            ("FR-06", "The user can draw with pen tools and erase strokes.", "High", "Yes"),
            ("FR-07", "The user can add sticky notes and emoji stickers to a note page.", "Medium", "Yes"),
            ("FR-08", "The user can create, complete, and delete tasks.", "High", "Yes"),
            ("FR-09", "The user can create and delete dated reminders.", "High", "Yes"),
            ("FR-10", "The user can search saved data and export a local backup.", "Medium", "Yes"),
        ],
        widths=[0.7, 4.3, 0.9, 0.9],
    )
    add_heading(doc, "3.2 Non-Functional Requirements", 2)
    add_table(
        doc,
        ["Category", "Requirement", "How it is addressed"],
        [
            ("Usability", "The app should be easy to navigate.", "Drawer/sidebar navigation, quick actions, empty states, and dashboard counters guide the user."),
            ("Performance", "Local operations should respond quickly.", "Database calls are asynchronous and data is scoped by field and user."),
            ("Reliability", "The app should preserve data and handle empty states.", "SQLite/web storage persists records; empty screens show clear actions."),
            ("Maintainability", "Code should be organized into layers.", "Screens, controllers, repositories, DAO classes, database adapters, models, and widgets are separated."),
            ("Data safety", "User data should not mix between accounts.", "Rows include userEmail and repositories apply current-user scoping."),
            ("Portability", "The app should be testable on web and native targets.", "A shared AppDatabase interface supports SQLite on native and SharedPreferences on web."),
        ],
        widths=[1.2, 2.3, 3.4],
    )
    add_heading(doc, "3.3 Use Cases", 2)
    add_table(
        doc,
        ["Use case", "Actor", "Precondition", "Main flow", "Expected result"],
        [
            ("UC-01: Create subject", "User", "User is logged in", "Tap Subjects, tap Add, enter title, choose cover, save.", "Subject appears on the notebook shelf."),
            ("UC-02: Create note", "User", "At least one subject exists", "Open subject, choose note type, write or draw, tap Save.", "Note is saved for the subject."),
            ("UC-03: Add visual elements", "User", "Note editor is open", "Tap Sticky or Sticker, place element on the page.", "Sticky note or emoji sticker appears and is saved."),
            ("UC-04: Manage task", "User", "User is logged in", "Open Tasks, add task, mark complete or delete.", "Task list updates."),
            ("UC-05: Search data", "User", "Saved data exists", "Open Search and type a keyword.", "Matching subjects, notes, tasks, and reminders are shown."),
            ("UC-06: Export backup", "User", "User is logged in", "Open Backup and copy JSON export.", "Current field data is exported without password data."),
        ],
        widths=[1.2, 0.8, 1.2, 2.5, 1.4],
    )

    add_heading(doc, "4. User Interface Design and Screens", 1)
    add_heading(doc, "4.1 UI Design Approach", 2)
    add_paragraph(
        doc,
        "The UI uses Material 3 components with a student-notebook visual direction. The design combines functional productivity screens with visual touches such as notebook covers, a shelf preview, sticky notes, emoji stickers, theme colors, and pen-style icons. Compact web/mobile layouts use an AppBar and drawer, while wider layouts use a sidebar. Empty states include clear action buttons so the user knows what to do next.",
    )
    add_heading(doc, "4.2 Navigation Flow", 2)
    add_paragraph(
        doc,
        "Startup -> Login/Sign Up -> Field Selection -> Home Dashboard -> Subjects/Notes/Tasks/Reminders/Search/Resources/Theme/Backup. From Subjects, the user opens a specific subject and then enters the Subject Note editor. After saving or returning, dashboard counts refresh.",
    )
    add_heading(doc, "4.3 Main Screenshots and Explanations", 2)
    add_screen_table(
        doc,
        [
            ("Home dashboard", "01_mobile_home.png", "Shows selected mode, counts, quick access, and notebook shelf.", "User can refresh, open drawer, or use quick action chips."),
            ("Drawer navigation", "02_mobile_drawer.png", "Provides access to Home, Subjects, Notes, Tasks, Reminders, Search, Resources, Backup, Theme, and Logout.", "User opens the drawer on compact screens and selects a destination."),
            ("Subjects", "03_subjects.png", "Displays created subjects as notebook cards.", "User can open, edit, delete, or add a subject."),
            ("Subject cover", "04_subject_cover.png", "Shows the custom notebook cover used for subject identity.", "User selects cover styles while adding or editing a subject."),
            ("Note editor", "05_note_editor.png", "Provides typed text, drawing canvas, pen types, colors, eraser, templates, and save action.", "User writes, draws, erases, changes tool, and saves."),
            ("Sticky note editor", "06_note_sticky.png", "Shows a sticky note attached to the page for visual reminders.", "User can move or delete sticky notes and stickers."),
        ],
    )

    add_heading(doc, "5. Technical Architecture", 1)
    add_heading(doc, "5.1 Overall Layered Architecture", 2)
    add_table(
        doc,
        ["Layer", "Main responsibility", "Examples"],
        [
            ("Presentation Layer", "Screens, widgets, navigation, forms, dialogs, and user interaction.", "home_screen.dart, subject_note_screen.dart, notebook_cover.dart"),
            ("Controller Layer", "Receives UI actions, performs validation, and calls repositories.", "auth_controller.dart, note_controller.dart, task_controller.dart"),
            ("Repository Layer", "Provides clean app-facing APIs and applies user scoping.", "subject_repository.dart, note_repository.dart, backup_repository.dart"),
            ("DAO Layer", "Contains table-specific insert, update, delete, select, and search operations.", "subject_dao.dart, note_dao.dart, task_dao.dart"),
            ("Database Layer", "Creates database adapters and stores persistent data.", "db_helper_io.dart, db_helper_web.dart, app_database.dart"),
            ("Model/Utility Layer", "Represents data objects and shared utilities.", "subject_model.dart, user_model.dart, password_hasher.dart"),
        ],
        widths=[1.3, 3.2, 2.4],
    )
    add_paragraph(
        doc,
        "This structure improves maintainability because UI code does not directly own SQL queries. A screen calls a controller, the controller calls a repository, the repository applies the current user scope, and the DAO performs the table operation through the database adapter.",
    )
    add_heading(doc, "5.2 Folder and Package Structure", 2)
    add_paragraph(
        doc,
        "lib/ contains controllers/, dao/, database/, models/, repositories/, screens/, utils/, and widgets/. test/ contains automated tests for password hashing, web database behavior, backup export behavior, and widget startup. docs/ contains report files and screenshots. scripts/ contains report-generation utilities.",
    )
    add_heading(doc, "5.3 Responsibilities of Each Layer", 2)
    add_table(
        doc,
        ["Layer / component", "Responsibility in this project", "Example class or file"],
        [
            ("Presentation Layer", "Renders UI and handles user interaction.", "HomeScreen, SubjectsScreen, SubjectNoteScreen"),
            ("State/Logic", "Holds screen state and calls controller methods.", "setState in screens; AppThemeController ValueNotifier"),
            ("Repository", "Hides DAO/database details and adds user scope.", "NoteRepository, SubjectRepository, BackupRepository"),
            ("DAO", "Runs table-specific CRUD and search queries.", "NoteDao, SubjectDao, ReminderDao"),
            ("SQLite Database", "Creates tables, handles migrations, and stores data persistently.", "DBHelper in db_helper_io.dart"),
            ("Model / Entity", "Maps app data to Dart objects or row maps.", "SubjectModel, UserModel"),
        ],
        widths=[1.5, 3.3, 2.1],
    )

    add_heading(doc, "6. Data Layer: SQLite, DAO, and Repository", 1)
    add_heading(doc, "6.1 SQLite Database Design", 2)
    add_table(
        doc,
        ["Table name", "Column", "Data type", "Constraint / purpose"],
        [
            ("users", "id", "INTEGER", "Primary key, auto-increment"),
            ("users", "name", "TEXT", "User display name"),
            ("users", "email", "TEXT", "Unique login identifier"),
            ("users", "password", "TEXT", "PBKDF2 password hash or legacy value upgraded on login"),
            ("subjects", "id", "INTEGER", "Primary key, auto-increment"),
            ("subjects", "title", "TEXT", "Subject notebook title"),
            ("subjects", "field", "TEXT", "Selected study field or specialization"),
            ("subjects", "userEmail", "TEXT", "Scopes subject data to the active account"),
            ("subjects", "coverColor", "INTEGER", "Stores selected cover color"),
            ("subjects", "coverPattern", "TEXT", "Stores selected cover pattern"),
            ("notes", "id", "INTEGER", "Primary key, auto-increment"),
            ("notes", "subject", "TEXT", "Subject name connected to the note"),
            ("notes", "field", "TEXT", "Selected field"),
            ("notes", "content", "TEXT", "Typed note content"),
            ("notes", "noteType", "TEXT", "Blank, lined, grid, or template note"),
            ("notes", "drawing", "TEXT", "JSON payload for strokes, sticky notes, and stickers"),
            ("notes", "userEmail", "TEXT", "Scopes notes to the active account"),
            ("tasks", "id/title/isCompleted/field/userEmail", "INTEGER/TEXT", "Task checklist data scoped by user and field"),
            ("reminders", "id/title/reminderDate/field/userEmail", "INTEGER/TEXT", "Dated reminder data scoped by user and field"),
        ],
        widths=[1.0, 1.7, 1.0, 3.2],
    )
    add_paragraph(
        doc,
        "SQLite is appropriate because My Notebook is a local-first productivity app with structured records. Users expect subjects, notes, tasks, and reminders to remain available after closing the app. The database version and migration helpers add missing columns such as userEmail, coverColor, coverPattern, drawing, noteType, and content when older local databases are opened.",
    )
    add_heading(doc, "6.2 DAO Design", 2)
    add_table(
        doc,
        ["DAO method", "Purpose", "SQL operation", "Called by"],
        [
            ("SubjectDao.findByField", "Loads subjects for field and user.", "SELECT with WHERE", "SubjectRepository"),
            ("SubjectDao.insertSubject", "Adds a subject and cover metadata.", "INSERT", "SubjectRepository"),
            ("SubjectDao.updateSubject", "Updates subject title/cover.", "UPDATE", "SubjectRepository"),
            ("NoteDao.findBySubjectAndField", "Loads one saved note.", "SELECT with LIMIT", "NoteRepository"),
            ("NoteDao.upsertNote", "Creates or updates note data.", "INSERT or UPDATE", "NoteRepository"),
            ("TaskDao.updateCompletion", "Toggles task completion.", "UPDATE", "TaskRepository"),
            ("ReminderDao.findByField", "Loads reminders ordered by date.", "SELECT with ORDER BY", "ReminderRepository"),
            ("UserDao.findByEmail", "Finds a user account for login.", "SELECT", "AuthRepository"),
        ],
        widths=[1.9, 2.4, 1.2, 1.5],
    )
    add_heading(doc, "6.3 Repository Design", 2)
    add_paragraph(
        doc,
        "Repositories act as the boundary between app logic and table access. They obtain the current account email through CurrentUserScope, prevent data access when there is no active user, and call DAO methods with field and userEmail filters. BackupRepository also builds an export payload from subjects, notes, tasks, and reminders without including the users table or password data.",
    )
    add_heading(doc, "6.4 Data Flow Example", 2)
    add_table(
        doc,
        ["Step", "Layer/component", "What happens"],
        [
            ("1", "UI screen", "The user opens SubjectNoteScreen, writes content, draws strokes, adds sticky notes, and taps Save."),
            ("2", "Controller", "NoteController receives subject, field, content, note type, and encoded drawing data."),
            ("3", "Repository", "NoteRepository retrieves the active user email and calls NoteDao."),
            ("4", "DAO", "NoteDao checks for an existing note and performs INSERT or UPDATE."),
            ("5", "Database", "SQLite stores the row with subject, field, drawing JSON, and userEmail."),
            ("6", "UI", "The screen shows a SnackBar and the note remains available when reopened."),
        ],
        widths=[0.6, 1.8, 4.5],
    )

    add_heading(doc, "7. Implementation Details", 1)
    add_heading(doc, "7.1 Flutter Widgets and Screens", 2)
    add_paragraph(
        doc,
        "Major screens include LoginScreen, SignUpScreen, FieldSelectionScreen, HomeScreen, SubjectsScreen, NotesScreen, SubjectNoteScreen, TasksScreen, RemindersScreen, SearchScreen, ResourcesScreen, ThemeScreen, and BackupScreen. Reusable UI includes NotebookCover and cover swatches. SubjectNoteScreen is the most interactive screen because it combines text input, drawing, erasing, pen type selection, templates, sticky notes, and stickers.",
    )
    add_heading(doc, "7.2 Navigation and Routing", 2)
    add_paragraph(
        doc,
        "The application uses Flutter Navigator with MaterialPageRoute. Data such as selectedField, subjectTitle, and noteType is passed through constructor arguments. Compact layouts use a Drawer from HomeScreen; wider layouts use a sidebar. After returning from data-changing screens, HomeScreen reloads dashboard counts.",
    )
    add_heading(doc, "7.3 State Management", 2)
    add_paragraph(
        doc,
        "The project uses StatefulWidget and setState for local screen state. This is suitable because the project has moderate complexity and most state belongs to specific screens. Theme color is managed with a ValueNotifier in AppThemeController so MaterialApp can rebuild when the user selects a new seed color.",
    )
    add_heading(doc, "7.4 Validation, Error Handling, and User Feedback", 2)
    add_table(
        doc,
        ["Situation", "How the app handles it", "User feedback"],
        [
            ("Empty login/signup fields", "Controllers/screens validate before continuing.", "Form messages or no action until fields are valid."),
            ("Duplicate or legacy login data", "AuthRepository checks email and supports password hash upgrade.", "Successful login or invalid credential feedback."),
            ("Empty subject/task/reminder input", "Add methods return early if text is empty.", "Dialog remains open or no invalid record is saved."),
            ("Empty data list", "Screens display friendly empty states.", "Action buttons guide the user to create data."),
            ("Save note", "Note data is encoded and upserted.", "SnackBar says Note saved."),
            ("Backup export", "Only active user and field data is exported.", "Copy action shows Backup copied."),
        ],
        widths=[1.7, 3.2, 2.0],
    )
    add_heading(doc, "7.5 Important Code Snippets", 2)
    add_code_block(
        doc,
        "Snippet 1: User-scoped note lookup",
        "lib/dao/note_dao.dart",
        "where: 'subject = ? AND field = ? AND userEmail = ?'\nwhereArgs: [subject, field, userEmail]",
        "This query pattern prevents notes from one account appearing in another account's workspace.",
    )
    add_code_block(
        doc,
        "Snippet 2: Repository applies active user scope",
        "lib/repositories/subject_repository.dart",
        "final userEmail = await userScope.email();\nif (userEmail == null) return [];\nreturn await subjectDao.findByField(field: field, userEmail: userEmail);",
        "The repository is responsible for applying the current account before DAO access.",
    )
    add_code_block(
        doc,
        "Snippet 3: Backup excludes account credentials",
        "lib/repositories/backup_repository.dart",
        "static const List<String> _tables = [\n  'subjects',\n  'notes',\n  'tasks',\n  'reminders',\n];",
        "The export intentionally includes notebook data only and does not include the users table.",
    )

    add_heading(doc, "8. Testing and Quality Assurance", 1)
    add_table(
        doc,
        ["Test ID", "Feature tested", "Test steps", "Expected result", "Actual result"],
        [
            ("T-01", "Password hashing", "Run password_hasher_test.dart.", "Passwords hash and verify correctly.", "Pass"),
            ("T-02", "Legacy password upgrade", "Run password_hasher_test.dart legacy case.", "Legacy plaintext can be upgraded safely.", "Pass"),
            ("T-03", "Web database filtering", "Run web_database_test.dart.", "Rows filter by field, user, and search pattern.", "Pass"),
            ("T-04", "Legacy row claim", "Run web_database_test.dart claim case.", "Unowned rows can be assigned to logged-in user.", "Pass"),
            ("T-05", "Backup export safety", "Run backup_repository_test.dart.", "Only current user's selected field is exported.", "Pass"),
            ("T-06", "Startup widget", "Run widget_test.dart.", "App shows login screen with no active session.", "Pass"),
            ("T-07", "Static analysis", "Run flutter analyze.", "No analyzer issues.", "Pass"),
            ("T-08", "Chrome launch", "Run flutter run -d chrome.", "App launches without ListTile/ColoredBox runtime error.", "Pass"),
        ],
        widths=[0.7, 1.6, 2.0, 1.7, 0.8],
    )
    add_paragraph(
        doc,
        "Manual quality checks were also performed for Chrome and macOS/native run commands. Earlier build checks passed for web and Android debug builds. During development, UI overflow on compact viewport was identified and fixed, and a runtime Material/ListTile issue in the sidebar was corrected by replacing a colored Container wrapper with a Material surface.",
    )

    add_heading(doc, "9. Challenges, Design Decisions, and Solutions", 1)
    add_heading(doc, "9.1 Technical Challenges Encountered", 2)
    add_table(
        doc,
        ["Challenge", "Cause / context", "Solution applied", "Result / lesson learned"],
        [
            ("Logged-in startup routing", "Initial startup could send users back to login.", "Startup now restores logged-in state and saved field.", "Better session continuity."),
            ("Plain-text passwords", "Early local storage used plain text passwords.", "Added PBKDF2 SHA-256 hashing and legacy upgrade support.", "Improved data safety."),
            ("User data separation", "Data was separated by field but not fully by account.", "Added userEmail scoping across subject, note, task, and reminder queries.", "Multiple users no longer share field data."),
            ("Chrome white screen/runtime issue", "Flutter web surfaced runtime/material problems during QC.", "Fixed sidebar Material wrapper and verified Chrome launch.", "Runtime UI errors must be tested, not only analyzed."),
            ("Basic note editor feel", "Original editor felt like a common note app.", "Added pen types, eraser size behavior, sticky notes, emoji stickers, and notebook covers.", "The app became more visual and distinctive."),
        ],
        widths=[1.3, 1.9, 2.2, 1.6],
    )
    add_heading(doc, "9.2 Design Decisions", 2)
    add_table(
        doc,
        ["Decision", "Alternatives considered", "Reason for final choice"],
        [
            ("Use SQLite for native persistence", "Only SharedPreferences, Firebase, in-memory list", "SQLite fits structured local records such as subjects, notes, tasks, and reminders."),
            ("Use Repository and DAO layers", "Direct database calls from screens", "Separation improves maintainability and matches the course architecture requirement."),
            ("Use setState for screen state", "Provider, Riverpod, Bloc", "The app's current state needs are screen-local and manageable without heavier state frameworks."),
            ("Use direct Navigator routes", "Named routes or go_router", "The project has simple route flow and passes a few constructor arguments."),
            ("Add visual notebook features", "Plain text-only notes", "Covers, pen tools, sticky notes, and stickers make the app more useful and less generic."),
        ],
        widths=[1.6, 2.0, 3.2],
    )
    add_heading(doc, "9.3 Known Limitations", 2)
    add_bullet(doc, "Cloud sync and multi-device backup are not included in this school-submission version.")
    add_bullet(doc, "Backup export is implemented, but backup import/restore is future work.")
    add_bullet(doc, "The app is not yet prepared for App Store or Google Play release signing.")
    add_bullet(doc, "More automated tests can be added for full UI flows such as subject creation and note editing.")
    add_bullet(doc, "Advanced note export such as PDF export is not yet implemented.")

    add_heading(doc, "10. Conclusion and Future Work", 1)
    add_heading(doc, "10.1 Conclusion", 2)
    add_paragraph(
        doc,
        "My Notebook meets its purpose as a student-focused Flutter notebook application. It provides meaningful academic workflows, user-scoped local data, visual note editing, task/reminder support, search, resources, theme customization, and backup export. The project also demonstrates course-relevant technical design through SQLite persistence, DAO classes, repository abstraction, controllers, reusable widgets, and automated tests.",
    )
    add_heading(doc, "10.2 Future Improvements", 2)
    add_table(
        doc,
        ["Improvement idea", "Why it is useful", "Possible implementation approach"],
        [
            ("Backup import/restore", "Allows users to recover exported data.", "Parse backup JSON and insert rows through repositories."),
            ("Cloud synchronization", "Supports multiple devices and safer backup.", "Firebase, Supabase, or a REST backend."),
            ("PDF export", "Lets users submit or share notes.", "Generate PDFs from note content and drawing data."),
            ("More templates", "Improves study workflows.", "Add Cornell notes, revision cards, diagrams, and planner pages."),
            ("Release preparation", "Needed for App Store/Google Play.", "Add icon, splash screen, privacy policy, signing, and release builds."),
        ],
        widths=[1.6, 2.3, 3.0],
    )

    add_heading(doc, "11. References", 1)
    add_table(
        doc,
        ["Source", "Type", "How it was used"],
        [
            ("Flutter documentation", "Official documentation", "Used for widgets, Material app structure, navigation, and platform builds."),
            ("Dart documentation", "Official documentation", "Used for language features and asynchronous programming."),
            ("sqflite package documentation", "Package documentation", "Used for SQLite local persistence on native targets."),
            ("shared_preferences package documentation", "Package documentation", "Used for sessions, theme color, and web adapter persistence."),
            ("Material Design 3 guidance", "Design system", "Used for UI components, theme color, cards, buttons, and navigation styling."),
            ("Flutter test documentation", "Official documentation", "Used for widget and unit tests."),
            ("AI assistance disclosure", "Tool support", "Used for coding support, debugging, report drafting, and project organization under student direction."),
        ],
        widths=[2.0, 1.5, 3.4],
    )

    add_heading(doc, "12. Appendices", 1)
    add_heading(doc, "Appendix A: Additional Screenshots", 2)
    add_screen_table(
        doc,
        [
            ("Mobile home", "01_mobile_home.png", "Responsive dashboard with counters and mode context.", "User reviews summary and opens navigation."),
            ("Mobile drawer", "02_mobile_drawer.png", "Compact navigation for major modules.", "User selects destination screen."),
            ("Sticky note page", "06_note_sticky.png", "Visual note page with sticky note element.", "User places movable notes on the page."),
        ],
    )
    add_heading(doc, "Appendix B: Repository and Submission Links", 2)
    add_table(
        doc,
        ["Item", "Link / note"],
        [
            ("Source code repository", "TO FILL: paste GitHub/GitLab/Drive link."),
            ("Demo video", "TO FILL: paste YouTube video link if required."),
            ("APK / release file", "TO FILL: paste APK or app bundle link if requested."),
            ("Other files", "Report file, screenshots, and project folder are included locally."),
        ],
        widths=[2.0, 4.9],
    )
    add_heading(doc, "Appendix C: Final Self-Check Before Submission", 2)
    for item in [
        "The report explains the project purpose, problem addressed, scope, target users, and core features.",
        "Main functional screens are shown with screenshots, captions, screen purposes, user actions, and navigation paths.",
        "The layered architecture is explained with UI, controllers, repositories, DAO, and SQLite responsibilities.",
        "SQLite tables, DAO methods, Repository methods, data flow, and state flow are explained.",
        "Technical challenges, causes, design decisions, solutions, tests, limitations, and future improvements are included.",
        "Testing evidence is included.",
        "References and external resources are acknowledged.",
        "Student number, signature, source code link, and demo video link are filled before submission.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
